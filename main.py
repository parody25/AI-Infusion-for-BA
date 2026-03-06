from fastapi import FastAPI, UploadFile, File, HTTPException, Body
from fastapi.responses import JSONResponse, FileResponse
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
from typing import Optional
import os
import uuid
import json
import tempfile
import shutil
from datetime import datetime
from dotenv import load_dotenv
from llama_parse import LlamaParse
from llama_index.llms.openai import OpenAI as LlamaOpenAI
from llama_index.embeddings.openai import OpenAIEmbedding
from llama_index.core import VectorStoreIndex, load_index_from_storage, StorageContext, Settings
from llama_index.core.node_parser import MarkdownElementNodeParser
from services.openai_brd_service import OpenAIBRDService
from services.openai_user_stories_service import OpenAIUserStoriesService
from services.jira_service import JiraService, get_jira_config, JiraConfig
from typing import List, Dict
from email import policy
from email.parser import BytesParser
from brd_schema import BRD_SCHEMA_JSON_STRING
from user_stories_schema import USER_STORIES_SCHEMA_JSON_STRING
from document_download_routes import doc_download_router

load_dotenv()

app = FastAPI(title="AI Infusion for BA", description="POC for enhancing BRD quality with AI")
app.include_router(doc_download_router)

# Add CORS middleware
app.add_middleware(
    CORSMiddleware,
    allow_origins=["http://localhost:3000"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Initialize clients
llama_parse = LlamaParse(api_key=os.getenv("LLAMA_CLOUD_API_KEY"))
llm = LlamaOpenAI(model="gpt-5.1", api_key=os.getenv("OPENAI_API_KEY"))
embed_model = OpenAIEmbedding(model="text-embedding-3-large", api_key=os.getenv("OPENAI_API_KEY"))

# CRITICAL FIX: Set global LlamaIndex settings to ensure consistent embedding model
Settings.embed_model = embed_model
Settings.llm = llm

brd_service = OpenAIBRDService()

# Directory for storing BRD projects
PROJECTS_DIR = "brd_projects"

# Dummy BRD Template with placeholders (First Islamic Bank format - matches JSON schema)
BRD_TEMPLATE = """
# First Islamic Bank - Business Requirements Specification

## Title: {title_main}
## ID: {id}
## Program: {program}
## Type: {type}

## Overview
{overview}

## Current Constraints
{current_constraint}

## Objective
{objective}

## In Scope
{in_scope}

## Out of Scope
{out_of_scope}

## Business Requirements
REQ ID: {req_id_bs}
Title: {title_bs}
Description: {description}
AS IS Behavior: {as_is_behaviour}
TO BE Behavior: {to_be_behaviour}
Pre-requisite: {pre_requisite}
Acceptance Criteria: {acceptance_criteria}
Alternate Flows: {alternate_flows}

## Reference Documents
{reference_documents}

## Requirement Traceability Matrix
REQ ID: {req_id}
Description: {description}
Source Channel: {source_channel}
Impacted System: {impacted_system}
Outcome: {outcome}

## Non-Functional Requirements
No. of users: {no_of_users}
Peak Volume: {peak_volume}
Monthly Volume: {monthly_volume}
Availability: {availability}
Impact on Operational Process: {impact_on_operational_process}
Regulatory Impact: {regulatory_impact}
Reports Requirement: {reports_requirement}
Access Requirement: {access_requirement}
Security Requirement: {secureity_requirement}
Data Requirement: {date_requirement}
Training Requirement: {training_requirement}
"""

def parse_eml(file_path: str) -> str:
    """
    Parse .eml file to extract clean text content from email bodies.
    Prefers text/plain over text/html, strips MIME headers and formatting.
    """
    try:
        with open(file_path, "rb") as f:
            msg = BytesParser(policy=policy.default).parse(f)

        parts = []

        for part in msg.walk():
            content_type = part.get_content_type()

            if content_type == "text/plain":
                text = part.get_payload(decode=True)
                if text:
                    charset = part.get_content_charset() or "utf-8"
                    parts.append(text.decode(charset, errors="ignore"))

            elif content_type == "text/html" and not parts:
                # Only use HTML if no plain text found
                html = part.get_payload(decode=True)
                if html:
                    charset = part.get_content_charset() or "utf-8"
                    # Basic HTML tag stripping for HTML fallback
                    import re
                    html_text = html.decode(charset, errors="ignore")
                    # Remove basic HTML tags
                    clean_text = re.sub(r'<[^>]+>', '', html_text)
                    # Decode HTML entities
                    import html
                    clean_text = html.unescape(clean_text)
                    parts.append(clean_text)

        return "\n".join(parts).strip()
    except Exception as e:
        print(f"ERROR: Failed to parse .eml file {file_path}: {e}")
        # Fallback to plain text reading
        try:
            with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                return f.read().strip()
        except:
            return ""

if not os.path.exists(PROJECTS_DIR):
    os.makedirs(PROJECTS_DIR)

def get_project_path(project_id: str) -> str:
    """Get the path for a BRD project directory."""
    return os.path.join(PROJECTS_DIR, project_id)

def load_project_metadata(project_id: str) -> Dict:
    """Load project metadata including document and BRD lists."""
    project_path = get_project_path(project_id)
    metadata_path = os.path.join(project_path, "metadata.json")
    if os.path.exists(metadata_path):
        with open(metadata_path, 'r') as f:
            metadata = json.load(f)
            # Ensure brds array exists for backward compatibility
            if "brds" not in metadata:
                metadata["brds"] = []
            return metadata
    return {"documents": [], "brds": []}

def save_project_metadata(project_id: str, metadata: Dict):
    """Save project metadata."""
    project_path = get_project_path(project_id)
    os.makedirs(project_path, exist_ok=True)
    metadata_path = os.path.join(project_path, "metadata.json")
    with open(metadata_path, 'w') as f:
        json.dump(metadata, f, indent=2)

@app.post("/create_project")
async def create_brd_project(name: str = None):
    """Create a new BRD project with optional name for easy reference."""
    project_id = str(uuid.uuid4())
    project_path = get_project_path(project_id)
    os.makedirs(project_path, exist_ok=True)

    # Initialize metadata with optional name
    metadata = {
        "documents": [],
        "name": name,
        "created_at": datetime.now().isoformat()
    }
    save_project_metadata(project_id, metadata)

    response = {
        "project_id": project_id,
        "message": "BRD project created successfully"
    }
    if name:
        response["name"] = name

    return JSONResponse(content=response)

@app.get("/projects")
async def list_projects():
    """List all BRD projects with their names and metadata."""
    if not os.path.exists(PROJECTS_DIR):
        return JSONResponse(content={"projects": []})

    projects = []
    for item in os.listdir(PROJECTS_DIR):
        project_path = os.path.join(PROJECTS_DIR, item)
        if os.path.isdir(project_path):
            metadata = load_project_metadata(item)
            projects.append({
                "project_id": item,
                "name": metadata.get("name"),
                "document_count": len(metadata.get("documents", [])),
                "brd_count": len(metadata.get("brds", [])),
                "created_at": metadata.get("created_at")
            })

    return JSONResponse(content={"projects": projects})

@app.post("/projects/{project_id}/upload")
async def upload_document_to_project(project_id: str, file: UploadFile = File(...)):
    """Upload a document to a specific BRD project using LlamaIndex for efficient embeddings."""
    if not file.filename:
        raise HTTPException(status_code=400, detail="No file provided")

    # Check if project exists
    project_path = get_project_path(project_id)
    if not os.path.exists(project_path):
        raise HTTPException(status_code=404, detail="BRD project not found")

    # Create temp file
    with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as temp_file:
        temp_file.write(await file.read())
        temp_file_path = temp_file.name

    try:
        file_name = os.path.basename(file.filename)
        document_name, file_extension = os.path.splitext(file_name)
        document_id = f"{project_id}_{document_name}"

        # Create project embeddings directory
        project_embedding_dir = os.path.join(project_path, "embeddings")
        os.makedirs(project_embedding_dir, exist_ok=True)

        embeddings_file = os.path.join(project_embedding_dir, f"{file_name}_embeddings")

        print(f"DEBUG: Processing {file_name} for project {project_id}")
        print(f"DEBUG: Embeddings file path: {embeddings_file}")

        if file_extension.lower() in ['.pdf', '.docx', '.xlsx', '.txt', '.eml']:
            try:
                # Check if embeddings exist and are compatible
                recreate_embeddings = True
                if os.path.exists(embeddings_file):
                    try:
                        # Try to load and test the existing index
                        storage_context = StorageContext.from_defaults(persist_dir=embeddings_file)
                        test_index = load_index_from_storage(storage_context)
                        # Test with a dummy query to check compatibility
                        test_retriever = test_index.as_retriever(similarity_top_k=1)
                        test_nodes = test_retriever.retrieve("test")
                        print(f"DEBUG: Existing embeddings for {file_name} are compatible, reusing")
                        recreate_embeddings = False
                    except Exception as e:
                        print(f"DEBUG: Existing embeddings incompatible ({str(e)}), will recreate")
                        import shutil
                        shutil.rmtree(embeddings_file)
                        recreate_embeddings = True
                else:
                    print(f"DEBUG: No existing embeddings found for {file_name}, creating new")

                if recreate_embeddings:
                    # Create new embeddings using LlamaIndex
                    print(f"DEBUG: Creating new embeddings for {file_name}")

                    # Special handling for .eml files - parse MIME content first
                    processing_file_path = temp_file_path
                    if file_extension.lower() == '.eml':
                        print(f"DEBUG: Parsing .eml file content for clean text extraction")
                        clean_text = parse_eml(temp_file_path)
                        if clean_text.strip():
                            # Create a temporary text file with clean content
                            import tempfile as tmp
                            with tmp.NamedTemporaryFile(mode='w', suffix='.txt', delete=False, encoding='utf-8') as clean_file:
                                clean_file.write(clean_text)
                                processing_file_path = clean_file.name
                            print(f"DEBUG: Created clean text file for .eml processing: {len(clean_text)} characters")
                        else:
                            print(f"WARNING: Failed to extract clean text from .eml file, using original")

                    # Use LlamaParse with markdown result type for better structure
                    llama_parser = LlamaParse(result_type="markdown", api_key=os.getenv("LLAMA_CLOUD_API_KEY"))
                    documents = llama_parser.load_data(processing_file_path)
                    print(f"DEBUG: Parsed {len(documents)} documents from {file_name}")

                    # Clean up temporary text file if created for .eml
                    if file_extension.lower() == '.eml' and processing_file_path != temp_file_path:
                        try:
                            os.unlink(processing_file_path)
                        except:
                            pass

                    # Use MarkdownElementNodeParser for better multimodal support
                    node_parser = MarkdownElementNodeParser(llm=llm, num_workers=4)  # Reduced workers for stability
                    nodes = node_parser.get_nodes_from_documents(documents)
                    base_nodes, objects = node_parser.get_nodes_and_objects(nodes)
                    print(f"DEBUG: Created {len(base_nodes)} base nodes and {len(objects)} objects")

                    # Create VectorStoreIndex with multimodal support
                    index = VectorStoreIndex(base_nodes + objects, embed_model=embed_model)
                    print(f"DEBUG: Created VectorStoreIndex")

                    # Persist embeddings
                    index.storage_context.persist(embeddings_file)
                    print(f"DEBUG: Persisted embeddings to {embeddings_file}")

                # Update project metadata
                project_metadata = load_project_metadata(project_id)
                doc_info = {
                    "id": str(uuid.uuid4()),
                    "filename": file_name,
                    "uploaded_at": datetime.now().isoformat(),
                    "type": "business_document",
                    "embedding_path": embeddings_file
                }
                project_metadata["documents"].append(doc_info)
                save_project_metadata(project_id, project_metadata)

                print(f"DEBUG: Successfully processed {file_name} for project {project_id}")
                return JSONResponse(content={
                    "message": "Document uploaded and processed successfully",
                    "document_id": doc_info["id"],
                    "embedding_created": True
                })

            except Exception as e:
                print(f"ERROR: Failed to create embeddings for {file_name}: {str(e)}")
                import traceback
                traceback.print_exc()
                raise HTTPException(status_code=500, detail=f"Error processing document: {str(e)}")
        else:
            raise HTTPException(status_code=400, detail="Unsupported file type. Only PDF, DOCX, XLSX, TXT, and EML are supported.")

    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error processing upload: {str(e)}")
    finally:
        # Clean up temp file
        if os.path.exists(temp_file_path):
            os.remove(temp_file_path)

@app.get("/projects/{project_id}/documents")
async def list_project_documents(project_id: str):
    """List all documents uploaded to a BRD project."""
    project_path = get_project_path(project_id)
    if not os.path.exists(project_path):
        raise HTTPException(status_code=404, detail="BRD project not found")

    metadata = load_project_metadata(project_id)
    return JSONResponse(content={"documents": metadata["documents"]})

@app.delete("/projects/{project_id}/documents/{document_id}")
async def delete_project_document(project_id: str, document_id: str):
    """Delete a document from a BRD project."""
    project_path = get_project_path(project_id)
    if not os.path.exists(project_path):
        raise HTTPException(status_code=404, detail="BRD project not found")

    metadata = load_project_metadata(project_id)

    # Find and remove document from metadata
    doc_to_remove = None
    for doc in metadata["documents"]:
        if doc["id"] == document_id:
            doc_to_remove = doc
            break

    if not doc_to_remove:
        raise HTTPException(status_code=404, detail="Document not found in project")

    # Note: For POC, we're not rebuilding the vectorstore after deletion
    # In production, you'd need to rebuild or mark chunks as deleted
    metadata["documents"].remove(doc_to_remove)
    save_project_metadata(project_id, metadata)

    return JSONResponse(content={"message": "Document deleted successfully"})

@app.delete("/projects/{project_id}")
async def delete_project(project_id: str):
    """Delete an entire BRD project and all its associated data."""
    project_path = get_project_path(project_id)
    if not os.path.exists(project_path):
        raise HTTPException(status_code=404, detail="BRD project not found")

    try:
        # Delete the entire project directory and all contents
        # This includes: metadata.json, embeddings/, brds/, all documents and BRDs
        import shutil
        shutil.rmtree(project_path)
        print(f"DEBUG: Successfully deleted project directory: {project_path}")

        return JSONResponse(content={"message": "Project deleted successfully"})

    except Exception as e:
        print(f"ERROR: Failed to delete project {project_id}: {e}")
        raise HTTPException(status_code=500, detail=f"Failed to delete project: {str(e)}")

@app.put("/projects/{project_id}")
async def rename_project(project_id: str, request_data: Dict[str, str] = Body(...)):
    """Rename a BRD project."""
    project_path = get_project_path(project_id)
    if not os.path.exists(project_path):
        raise HTTPException(status_code=404, detail="BRD project not found")

    new_name = request_data.get("name", "").strip()
    if not new_name:
        raise HTTPException(status_code=400, detail="Project name cannot be empty")

    try:
        # Load current metadata
        metadata = load_project_metadata(project_id)

        # Update the name
        metadata["name"] = new_name

        # Save updated metadata
        save_project_metadata(project_id, metadata)

        print(f"DEBUG: Successfully renamed project {project_id} to '{new_name}'")

        return JSONResponse(content={"message": "Project renamed successfully"})

    except Exception as e:
        print(f"ERROR: Failed to rename project {project_id}: {e}")
        raise HTTPException(status_code=500, detail=f"Failed to rename project: {str(e)}")

class BRDGenerationRequest(BaseModel):
    requirements: str
    process_instructions: Optional[str] = None

@app.post("/projects/{project_id}/generate_brd")
async def generate_brd_for_project(project_id: str, request_data: BRDGenerationRequest):
    """Generate BRD for a specific project, store it in the project, and return metadata."""
    requirements = request_data.requirements
    process_instructions = request_data.process_instructions or ""
    
    print(f"DEBUG: Starting BRD generation for project {project_id}")
    print(f"DEBUG: Requirements: {requirements}")
    print(f"DEBUG: Process Instructions: {process_instructions}")

    project_path = get_project_path(project_id)
    if not os.path.exists(project_path):
        raise HTTPException(status_code=404, detail="BRD project not found")

    # Check if project has documents
    metadata = load_project_metadata(project_id)
    if not metadata["documents"]:
        raise HTTPException(status_code=400, detail="No documents uploaded to this project yet")

    print(f"DEBUG: Found {len(metadata['documents'])} documents in project")

    # Collect all embedding paths and document IDs for this project
    embedding_paths = []
    input_document_ids = []
    for doc in metadata["documents"]:
        embedding_path = doc.get("embedding_path")
        if embedding_path and os.path.exists(embedding_path):
            embedding_paths.append(embedding_path)
            input_document_ids.append(doc["id"])

    print(f"DEBUG: Found {len(embedding_paths)} valid embedding paths")
    print(f"DEBUG: Input document IDs: {input_document_ids}")

    if not embedding_paths:
        raise HTTPException(status_code=500, detail="No valid embeddings found for project")

    # Load and combine indices from all documents
    combined_nodes = []
    for embedding_path in embedding_paths:
        try:
            print(f"DEBUG: Loading index from {embedding_path}")
            storage_context = StorageContext.from_defaults(persist_dir=embedding_path)
            index = load_index_from_storage(storage_context)
            print(f"DEBUG: Successfully loaded index")

            # Get nodes from index
            retriever = index.as_retriever(similarity_top_k=15)
            nodes = retriever.retrieve(requirements)
            print(f"DEBUG: Retrieved {len(nodes)} nodes from {embedding_path}")
            combined_nodes.extend(nodes)
        except Exception as e:
            print(f"ERROR: Failed to load/process index from {embedding_path}: {e}")
            import traceback
            traceback.print_exc()
            continue

    print(f"DEBUG: Total combined nodes: {len(combined_nodes)}")

    # Extract context from retrieved nodes
    context_parts = []
    for i, node in enumerate(combined_nodes[:10]):  # Limit to top 10
        if hasattr(node, 'text') and node.text:
            context_parts.append(f"Context {i+1}: {node.text[:500]}...")  # Truncate for logging
        elif hasattr(node, 'content') and node.content:
            context_parts.append(f"Context {i+1}: {node.content[:500]}...")

    context = "\n\n".join(context_parts)
    print(f"DEBUG: Final context length: {len(context)} characters")

    if not context.strip():
        context = "No relevant context found in uploaded documents."
        print("WARNING: No context found, using default message")

    print("DEBUG: Calling BRD service to generate content")

    # Create BRDs directory in project
    brds_dir = os.path.join(project_path, "brds")
    os.makedirs(brds_dir, exist_ok=True)

    # Generate unique BRD ID and filename
    brd_id = str(uuid.uuid4())
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    output_filename = f"BRD_{timestamp}_{brd_id[:8]}.docx"
    output_path = os.path.join(brds_dir, output_filename)

    # Generate and fill Word document with BA instructions
    filled_path = brd_service.generate_brd_word(
        requirements=requirements,
        context=context,
        schema_json=BRD_SCHEMA_JSON_STRING,
        template_path="BRD_Template.docx",
        output_path=output_path,
        process_instructions=process_instructions
    )

    print(f"DEBUG: BRD generated successfully at {filled_path}")

    # Record BRD in project metadata
    brd_info = {
        "id": brd_id,
        "filename": output_filename,
        "file_path": output_path,
        "generated_at": datetime.now().isoformat(),
        "requirements": requirements,
        "process_instructions": process_instructions,
        "input_document_ids": input_document_ids,
        "document_count": len(input_document_ids)
    }

    metadata["brds"].append(brd_info)
    save_project_metadata(project_id, metadata)

    print(f"DEBUG: BRD recorded in metadata: {brd_id}")

    return JSONResponse(content={
        "message": "BRD generated and stored successfully",
        "brd_id": brd_id,
        "filename": output_filename,
        "input_documents_used": len(input_document_ids),
        "generated_at": brd_info["generated_at"]
    })

@app.get("/projects/{project_id}/brds")
async def list_project_brds(project_id: str):
    """List all BRDs generated for a project."""
    project_path = get_project_path(project_id)
    if not os.path.exists(project_path):
        raise HTTPException(status_code=404, detail="BRD project not found")

    metadata = load_project_metadata(project_id)
    brds = metadata.get("brds", [])

    # Add document filenames for better UX
    documents = metadata.get("documents", [])
    doc_id_to_filename = {doc["id"]: doc["filename"] for doc in documents}

    for brd in brds:
        input_doc_filenames = []
        for doc_id in brd.get("input_document_ids", []):
            filename = doc_id_to_filename.get(doc_id, f"Unknown (ID: {doc_id})")
            input_doc_filenames.append(filename)
        brd["input_document_filenames"] = input_doc_filenames

    return JSONResponse(content={"brds": brds})

@app.get("/projects/{project_id}/brds/{brd_id}/download")
async def download_project_brd(project_id: str, brd_id: str):
    """Download a specific BRD from a project."""
    project_path = get_project_path(project_id)
    if not os.path.exists(project_path):
        raise HTTPException(status_code=404, detail="BRD project not found")

    metadata = load_project_metadata(project_id)
    brds = metadata.get("brds", [])

    # Find the BRD
    brd_info = None
    for brd in brds:
        if brd["id"] == brd_id:
            brd_info = brd
            break

    if not brd_info:
        raise HTTPException(status_code=404, detail="BRD not found in project")

    file_path = brd_info["file_path"]
    if not os.path.exists(file_path):
        raise HTTPException(status_code=404, detail="BRD file not found on disk")

    return FileResponse(
        path=file_path,
        filename=brd_info["filename"],
        media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )

@app.delete("/projects/{project_id}/brds/{brd_id}")
async def delete_project_brd(project_id: str, brd_id: str):
    """Delete a BRD from a project."""
    project_path = get_project_path(project_id)
    if not os.path.exists(project_path):
        raise HTTPException(status_code=404, detail="BRD project not found")

    metadata = load_project_metadata(project_id)
    brds = metadata.get("brds", [])

    # Find and remove BRD from metadata
    brd_to_remove = None
    for brd in brds:
        if brd["id"] == brd_id:
            brd_to_remove = brd
            break

    if not brd_to_remove:
        raise HTTPException(status_code=404, detail="BRD not found in project")

    # Remove the file if it exists
    file_path = brd_to_remove.get("file_path")
    if file_path and os.path.exists(file_path):
        os.remove(file_path)

    # Remove from metadata
    metadata["brds"].remove(brd_to_remove)
    save_project_metadata(project_id, metadata)

    return JSONResponse(content={"message": "BRD deleted successfully"})

@app.get("/brd_template")
async def get_brd_template():
    """Get the BRD template structure."""
    return JSONResponse(content={"template": BRD_TEMPLATE})

class UserStoriesGenerationRequest(BaseModel):
    brd_id: str
    version: str

@app.post("/projects/{project_id}/generate_user_stories")
async def generate_user_stories_for_project(project_id: str, request_data: UserStoriesGenerationRequest):
    """Generate User Stories for a specific BRD version, store it in the project, and return metadata."""
    brd_id = request_data.brd_id
    version = request_data.version
    
    print(f"DEBUG: Starting User Stories generation for project {project_id}")
    print(f"DEBUG: BRD ID: {brd_id}")
    print(f"DEBUG: Version: {version}")

    project_path = get_project_path(project_id)
    if not os.path.exists(project_path):
        raise HTTPException(status_code=404, detail="BRD project not found")

    # Check if project has BRDs
    metadata = load_project_metadata(project_id)
    brds = metadata.get("brds", [])

    # Find the specific BRD
    brd_info = None
    for brd in brds:
        if brd["id"] == brd_id:
            brd_info = brd
            break

    if not brd_info:
        raise HTTPException(status_code=404, detail="BRD not found in project")

    # Read the BRD content
    file_path = brd_info["file_path"]
    if not os.path.exists(file_path):
        raise HTTPException(status_code=404, detail="BRD file not found on disk")

    # Extract text content from the BRD document
    try:
        from docx import Document
        doc = Document(file_path)
        brd_content = []
        for para in doc.paragraphs:
            if para.text.strip():
                brd_content.append(para.text)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        if para.text.strip():
                            brd_content.append(para.text)
        
        brd_text = "\n".join(brd_content)
        print(f"DEBUG: Extracted BRD content length: {len(brd_text)} characters")
    except Exception as e:
        print(f"ERROR: Failed to extract text from BRD: {e}")
        raise HTTPException(status_code=500, detail=f"Error reading BRD file: {str(e)}")

    if not brd_text.strip():
        raise HTTPException(status_code=400, detail="BRD content is empty")

    print("DEBUG: Calling User Stories service to generate content")

    # Create User Stories directory in project
    user_stories_dir = os.path.join(project_path, "user_stories")
    os.makedirs(user_stories_dir, exist_ok=True)

    # Generate unique User Stories ID and filename
    user_stories_id = str(uuid.uuid4())
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    output_filename = f"User_Stories_{timestamp}_{user_stories_id[:8]}.xlsx"
    output_path = os.path.join(user_stories_dir, output_filename)

    # Initialize User Stories service
    user_stories_service = OpenAIUserStoriesService()

    # Create BRD embeddings for RAG pipeline if they don't exist
    brd_embeddings_dir = os.path.join(project_path, "brd_embeddings")
    os.makedirs(brd_embeddings_dir, exist_ok=True)
    brd_embeddings_file = os.path.join(brd_embeddings_dir, f"{brd_id}_embeddings")

    print(f"DEBUG: Checking for existing BRD embeddings at {brd_embeddings_file}")
    
    # Create embeddings if they don't exist
    if not os.path.exists(brd_embeddings_file):
        print(f"DEBUG: Creating new BRD embeddings for {brd_info['filename']}")
        success = user_stories_service.create_brd_embeddings(file_path, brd_embeddings_file)
        if not success:
            print("WARNING: Failed to create BRD embeddings, will use direct text extraction")
    else:
        print(f"DEBUG: Using existing BRD embeddings for {brd_info['filename']}")

    # Retrieve context using RAG if embeddings exist
    if os.path.exists(brd_embeddings_file):
        print("DEBUG: Using RAG to retrieve relevant BRD context")
        rag_context = user_stories_service.retrieve_brd_context(
            brd_embeddings_file, 
            f"Generate user stories for version {version}", 
            top_k=30
        )
        
        if rag_context.strip():
            print(f"DEBUG: RAG context length: {len(rag_context)} characters")
            # Combine RAG context with direct text extraction for comprehensive coverage
            combined_content = f"=== RAG RETRIEVED CONTEXT ===\n{rag_context}\n\n=== FULL BRD CONTENT ===\n{brd_text}"
        else:
            print("DEBUG: RAG context empty, using direct text extraction only")
            combined_content = brd_text
    else:
        print("DEBUG: No embeddings available, using direct text extraction only")
        combined_content = brd_text

    # Generate and export User Stories to Excel
    try:
        filled_path = user_stories_service.generate_user_stories_excel(
            brd_content=combined_content,
            version=version,
            schema_json=USER_STORIES_SCHEMA_JSON_STRING,
            output_path=output_path
        )
        print(f"DEBUG: User Stories generated successfully at {filled_path}")
    except Exception as e:
        print(f"ERROR: Failed to generate User Stories: {e}")
        raise HTTPException(status_code=500, detail=f"Error generating User Stories: {str(e)}")

    # Record User Stories in project metadata
    user_stories_info = {
        "id": user_stories_id,
        "filename": output_filename,
        "file_path": output_path,
        "generated_at": datetime.now().isoformat(),
        "brd_id": brd_id,
        "version": version,
        "brd_filename": brd_info["filename"]
    }

    # Ensure user_stories array exists in metadata
    if "user_stories" not in metadata:
        metadata["user_stories"] = []
    
    metadata["user_stories"].append(user_stories_info)
    save_project_metadata(project_id, metadata)

    print(f"DEBUG: User Stories recorded in metadata: {user_stories_id}")

    return JSONResponse(content={
        "message": "User Stories generated and stored successfully",
        "user_stories_id": user_stories_id,
        "filename": output_filename,
        "brd_filename": brd_info["filename"],
        "version": version,
        "generated_at": user_stories_info["generated_at"]
    })

@app.get("/projects/{project_id}/user_stories")
async def list_project_user_stories(project_id: str):
    """List all User Stories generated for a project."""
    project_path = get_project_path(project_id)
    if not os.path.exists(project_path):
        raise HTTPException(status_code=404, detail="BRD project not found")

    metadata = load_project_metadata(project_id)
    user_stories = metadata.get("user_stories", [])

    # Add BRD filenames for better UX
    brds = metadata.get("brds", [])
    brd_id_to_filename = {brd["id"]: brd["filename"] for brd in brds}

    for us in user_stories:
        brd_filename = brd_id_to_filename.get(us.get("brd_id"), f"Unknown BRD (ID: {us.get('brd_id')})")
        us["brd_filename"] = brd_filename

    return JSONResponse(content={"user_stories": user_stories})

@app.get("/projects/{project_id}/user_stories/{user_stories_id}/download")
async def download_project_user_stories(project_id: str, user_stories_id: str):
    """Download a specific User Stories Excel file from a project."""
    project_path = get_project_path(project_id)
    if not os.path.exists(project_path):
        raise HTTPException(status_code=404, detail="BRD project not found")

    metadata = load_project_metadata(project_id)
    user_stories = metadata.get("user_stories", [])

    # Find the User Stories
    user_stories_info = None
    for us in user_stories:
        if us["id"] == user_stories_id:
            user_stories_info = us
            break

    if not user_stories_info:
        raise HTTPException(status_code=404, detail="User Stories not found in project")

    file_path = user_stories_info["file_path"]
    if not os.path.exists(file_path):
        raise HTTPException(status_code=404, detail="User Stories file not found on disk")

    return FileResponse(
        path=file_path,
        filename=user_stories_info["filename"],
        media_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
    )

@app.delete("/projects/{project_id}/user_stories/{user_stories_id}")
async def delete_project_user_stories(project_id: str, user_stories_id: str):
    """Delete a User Stories file from a project."""
    project_path = get_project_path(project_id)
    if not os.path.exists(project_path):
        raise HTTPException(status_code=404, detail="BRD project not found")

    metadata = load_project_metadata(project_id)
    user_stories = metadata.get("user_stories", [])

    # Find and remove User Stories from metadata
    us_to_remove = None
    for us in user_stories:
        if us["id"] == user_stories_id:
            us_to_remove = us
            break

    if not us_to_remove:
        raise HTTPException(status_code=404, detail="User Stories not found in project")

    # Remove the file if it exists
    file_path = us_to_remove.get("file_path")
    if file_path and os.path.exists(file_path):
        os.remove(file_path)

    # Remove from metadata
    metadata["user_stories"].remove(us_to_remove)
    save_project_metadata(project_id, metadata)

    return JSONResponse(content={"message": "User Stories deleted successfully"})

@app.get("/health")
async def health_check():
    return {"status": "healthy"}

# Jira Integration Endpoints

@app.get("/jira/config")
async def get_jira_config_status():
    """Get current Jira configuration status."""
    config = get_jira_config()
    if config:
        return JSONResponse(content={
            "configured": True,
            "server": config.server,
            "email": config.email,
            "default_project": config.default_project
        })
    else:
        return JSONResponse(content={
            "configured": False,
            "message": "Jira is not configured. Please set JIRA_SERVER, JIRA_EMAIL, and JIRA_API_TOKEN environment variables."
        })

@app.post("/jira/test-connection")
async def test_jira_connection():
    """Test Jira connection with current configuration."""
    config = get_jira_config()
    if not config:
        raise HTTPException(status_code=400, detail="Jira is not configured")
    
    jira_service = JiraService(config)
    result = jira_service.test_connection()
    
    if result["success"]:
        return JSONResponse(content=result)
    else:
        raise HTTPException(status_code=500, detail=result["message"])

@app.get("/jira/projects")
async def get_jira_projects():
    """Get list of available Jira projects."""
    config = get_jira_config()
    if not config:
        raise HTTPException(status_code=400, detail="Jira is not configured")
    
    jira_service = JiraService(config)
    projects = jira_service.get_projects()
    
    if not projects:
        raise HTTPException(status_code=500, detail="Failed to retrieve projects")
    
    return JSONResponse(content={"projects": projects})

@app.get("/jira/issue-types")
async def get_jira_issue_types():
    """Get list of available Jira issue types."""
    config = get_jira_config()
    if not config:
        raise HTTPException(status_code=400, detail="Jira is not configured")
    
    jira_service = JiraService(config)
    issue_types = jira_service.get_issue_types()
    
    if not issue_types:
        raise HTTPException(status_code=500, detail="Failed to retrieve issue types")
    
    return JSONResponse(content={"issue_types": issue_types})

class JiraSyncRequest(BaseModel):
    jira_url: Optional[str] = None
    project_key: Optional[str] = None
    auth_token: Optional[str] = None

@app.post("/projects/{project_id}/user_stories/{user_stories_id}/jira-sync")
async def sync_user_stories_to_jira(project_id: str, user_stories_id: str, request_data: JiraSyncRequest):
    """Sync User Stories to Jira."""
    print(f"🔍 DEBUG: Received sync request for project {project_id}, user stories {user_stories_id}")
    print(f"🔍 DEBUG: Request data: {request_data}")
    
    # Check if project exists
    project_path = get_project_path(project_id)
    if not os.path.exists(project_path):
        print(f"🔍 DEBUG: Project path {project_path} does not exist")
        raise HTTPException(status_code=404, detail="BRD project not found")

    # Load project metadata
    metadata = load_project_metadata(project_id)
    user_stories = metadata.get("user_stories", [])
    print(f"🔍 DEBUG: Found {len(user_stories)} user stories in project")

    # Find the User Stories
    user_stories_info = None
    for us in user_stories:
        if us["id"] == user_stories_id:
            user_stories_info = us
            break

    if not user_stories_info:
        print(f"🔍 DEBUG: User Stories {user_stories_id} not found in project")
        raise HTTPException(status_code=404, detail="User Stories not found in project")

    file_path = user_stories_info["file_path"]
    print(f"🔍 DEBUG: User Stories file path: {file_path}")
    if not os.path.exists(file_path):
        print(f"🔍 DEBUG: User Stories file does not exist at {file_path}")
        raise HTTPException(status_code=404, detail="User Stories file not found on disk")

    # Load User Stories data from Excel
    try:
        import pandas as pd
        excel_data = pd.read_excel(file_path, sheet_name=None)
        
        # Process User Stories sheet
        stories_data = []
        if 'User Stories' in excel_data:
            stories_df = excel_data['User Stories']
            for _, row in stories_df.iterrows():
                stories_data.append({
                    "type": "user_story",
                    "story_id": row.get('story_id', ''),
                    "title": row.get('title', ''),
                    "user_role": row.get('user_role', ''),
                    "description": row.get('description', ''),
                    "acceptance_criteria": row.get('acceptance_criteria', ''),
                    "priority": row.get('priority', 'Medium'),
                    "effort_estimate": row.get('effort_estimate', None),
                    "brd_reference": row.get('brd_reference', ''),
                    "version": row.get('version', '')
                })

        # Process Epics sheet
        epics_data = []
        if 'Epics' in excel_data:
            epics_df = excel_data['Epics']
            for _, row in epics_df.iterrows():
                epics_data.append({
                    "type": "epic",
                    "epic_id": row.get('epic_id', ''),
                    "title": row.get('title', ''),
                    "description": row.get('description', ''),
                    "related_stories": row.get('related_stories', [])
                })

        # Process Dependencies sheet
        dependencies_data = []
        if 'Dependencies' in excel_data:
            deps_df = excel_data['Dependencies']
            for _, row in deps_df.iterrows():
                dependencies_data.append({
                    "type": "dependency",
                    "story_id": row.get('story_id', ''),
                    "depends_on": row.get('depends_on', ''),
                    "dependency_type": row.get('dependency_type', 'Blocks')
                })

        all_issues = stories_data + epics_data + dependencies_data
        
    except Exception as e:
        print(f"ERROR: Failed to read User Stories Excel file: {e}")
        raise HTTPException(status_code=500, detail=f"Error reading User Stories file: {str(e)}")

    # Initialize Jira service
    jira_config = get_jira_config()
    
    # Override config if provided in request
    if request_data.jira_url and request_data.project_key and request_data.auth_token:
        # Use environment email if available, otherwise use a placeholder
        env_email = os.getenv("JIRA_EMAIL", "")
        frontend_email = env_email if env_email else "temp@temp.com"
        
        jira_config = JiraConfig(
            server=request_data.jira_url,
            email=frontend_email,
            api_token=request_data.auth_token,
            default_project=request_data.project_key
        )
    elif not jira_config:
        raise HTTPException(status_code=400, detail="Jira is not configured and no credentials provided")

    jira_service = JiraService(jira_config)
    
    # Test connection
    connection_result = jira_service.test_connection()
    if not connection_result["success"]:
        raise HTTPException(status_code=500, detail=connection_result["message"])

    # Use project key from request or config
    project_key = request_data.project_key or jira_config.default_project
    if not project_key:
        raise HTTPException(status_code=400, detail="Project key not specified")

    # Bulk create issues
    try:
        print(f"🔍 DEBUG: Starting bulk create with {len(all_issues)} issues")
        print(f"🔍 DEBUG: Project key: {project_key}")
        print(f"🔍 DEBUG: Issues to create: {[issue.get('type', 'unknown') + ': ' + issue.get('title', 'no title') for issue in all_issues]}")
        
        results = jira_service.bulk_create_issues(project_key, all_issues)
        
        print(f"🔍 DEBUG: Bulk create completed")
        print(f"🔍 DEBUG: Success: {len(results.get('success', []))}, Failed: {len(results.get('failed', []))}")
        
        # Update User Stories metadata with sync status
        user_stories_info["jira_sync_status"] = "completed"
        user_stories_info["jira_sync_at"] = datetime.now().isoformat()
        user_stories_info["jira_sync_results"] = results
        save_project_metadata(project_id, metadata)

        return JSONResponse(content={
            "message": "User Stories synced to Jira successfully",
            "sync_results": results,
            "project_key": project_key,
            "user_stories_id": user_stories_id
        })
        
    except Exception as e:
        print(f"ERROR: Failed to sync to Jira: {e}")
        import traceback
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=f"Error syncing to Jira: {str(e)}")

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000)
