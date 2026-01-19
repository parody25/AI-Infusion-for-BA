from docx import Document
from docx.shared import Inches

def create_brd_template():
    doc = Document()

    # Title
    title = doc.add_heading('Business Requirements Document (BRD)', 0)

    # Introduction section
    doc.add_heading('1. Introduction', level=1)
    intro_para = doc.add_paragraph()
    intro_para.add_run('{introduction}')

    # Business Objectives
    doc.add_heading('2. Business Objectives', level=1)
    obj_para = doc.add_paragraph()
    obj_para.add_run('{business_objectives}')

    # Scope
    doc.add_heading('3. Scope', level=1)

    doc.add_heading('In-Scope', level=2)
    in_scope_para = doc.add_paragraph()
    in_scope_para.add_run('{in_scope}')

    doc.add_heading('Out-of-Scope', level=2)
    out_scope_para = doc.add_paragraph()
    out_scope_para.add_run('{out_of_scope}')

    # Business Requirements
    doc.add_heading('4. Business Requirements', level=1)
    biz_req_para = doc.add_paragraph()
    biz_req_para.add_run('{business_requirements}')

    # Functional Requirements
    doc.add_heading('5. Functional Requirements', level=1)
    func_req_para = doc.add_paragraph()
    func_req_para.add_run('{functional_requirements}')

    # Non-Functional Requirements
    doc.add_heading('6. Non-Functional Requirements', level=1)
    non_func_para = doc.add_paragraph()
    non_func_para.add_run('{non_functional_requirements}')

    # Assumptions & Constraints
    doc.add_heading('7. Assumptions & Constraints', level=1)
    assumptions_para = doc.add_paragraph()
    assumptions_para.add_run('{assumptions_constraints}')

    # Risks & Dependencies
    doc.add_heading('8. Risks & Dependencies', level=1)
    risks_para = doc.add_paragraph()
    risks_para.add_run('{risks_dependencies}')

    # Acceptance Criteria
    doc.add_heading('9. Acceptance Criteria', level=1)
    acceptance_para = doc.add_paragraph()
    acceptance_para.add_run('{acceptance_criteria}')

    # Open Questions
    doc.add_heading('10. Open Questions', level=1)
    questions_para = doc.add_paragraph()
    questions_para.add_run('{open_questions}')

    # Save the template
    doc.save('BRD_Template.docx')
    print("BRD Template created successfully!")

if __name__ == "__main__":
    create_brd_template()
