import os
import json
import requests
import base64
import time
from datetime import datetime
from typing import Dict
from copy import deepcopy
from openai import OpenAI
from dotenv import load_dotenv
from docx import Document
from docx.table import Table
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from PIL import Image

load_dotenv()


class OpenAIBRDService:
    """
    Enterprise-grade BRD generator.
    Pipeline:
    Documents → Embeddings → GPT-5.1 (JSON) → Word Template
    """

    def __init__(
        self,
        model: str = "gpt-5.1",
        max_output_tokens: int = 14000,
        temperature: float = 0.2,
        timeout: int = 600
    ):
        api_key = os.getenv("OPENAI_API_KEY")
        if not api_key:
            raise RuntimeError("OPENAI_API_KEY not configured")

        self.client = OpenAI(api_key=api_key)
        self.model = model
        self.max_output_tokens = max_output_tokens
        self.temperature = temperature
        self.timeout = timeout

    # ------------------------------------------------------------------
    # GPT RESPONSE HANDLING
    # ------------------------------------------------------------------

    def _extract_text(self, response) -> str:
        """
        Extract text from OpenAI Responses API response object.
        Supports both .output_text and structured .output[].content[] shapes.
        """
        if hasattr(response, "output_text") and response.output_text:
            return response.output_text.strip()

        texts = []
        for item in getattr(response, "output", []):
            for content in getattr(item, "content", []):
                if content.get("type") == "output_text":
                    texts.append(content.get("text", ""))

        return "\n".join(texts).strip()

    # ------------------------------------------------------------------
    # PROMPT
    # ------------------------------------------------------------------

    def _system_prompt(self) -> str:
        """Load system prompt from external file."""
        prompt_file = os.path.join(os.path.dirname(__file__), '..', 'prompts', 'system_prompt.txt')
        try:
            with open(prompt_file, 'r', encoding='utf-8') as f:
                return f.read().strip()
        except FileNotFoundError:
            return (
                "You are a Senior Business Analyst with deep BFSI experience.\n"
                "You produce regulatory-compliant Business Requirement Specifications.\n"
                "You respond ONLY in valid JSON following the provided schema."
            )

    # ------------------------------------------------------------------
    # JSON GENERATION
    # ------------------------------------------------------------------

    def _load_user_prompt_template(self) -> str:
        """Load user prompt template from external file."""
        template_file = os.path.join(os.path.dirname(__file__), '..', 'prompts', 'user_prompt_template.txt')
        try:
            with open(template_file, 'r', encoding='utf-8') as f:
                return f.read().strip()
        except FileNotFoundError:
            return """
You MUST return ONLY valid JSON that follows the structure of the JSON SCHEMA provided below.

CRITICAL RULES:
- The "description" fields in the JSON SCHEMA are your specific instructions for what to generate for that key.
- Do NOT include the "description" keys in your JSON output; only return the data keys and their generated values.
- No markdown, no explanations.
- Use BFSI / CBUAE terminology.
- Use clear "shall" statements
- Do NOT invent information
- Format all text content using dot bullet points (starting with "• ") instead of paragraphs.
- For process descriptions, use numbered lists (starting with "1. ") to show sequence

=== REQUIREMENTS ===
{requirements}

=== CONTEXT ===
{context}

=== BA PROCESS INSTRUCTIONS ===
{process_instructions}

If "BA PROCESS INSTRUCTIONS" are provided, generate a corresponding Mermaid.js diagram code in the 'process_visuals' section of the JSON. Use 'graph TD' for flowcharts or 'sequenceDiagram' for system interactions.

=== JSON SCHEMA ===
{schema_json}
"""

    def generate_brd_json(self, requirements: str, context: str, schema_json: str, process_instructions: str = "") -> Dict:
        user_prompt = self._load_user_prompt_template().format(
            requirements=requirements,
            context=context,
            process_instructions=process_instructions,
            schema_json=schema_json
        )

        response = self.client.responses.create(
            model=self.model,
            input=[
                {"role": "system", "content": self._system_prompt()},
                {"role": "user", "content": user_prompt}
            ],
            max_output_tokens=self.max_output_tokens,
            timeout=self.timeout
        )

        raw = self._extract_text(response)

        if not raw:
            raise RuntimeError("GPT returned empty output")

        try:
            data = json.loads(raw)
        except json.JSONDecodeError as e:
            raise RuntimeError(f"Invalid JSON from GPT: {e}\n\n{raw}")

        os.makedirs("llm_responses", exist_ok=True)
        with open(
            f"llm_responses/brd_{datetime.now().strftime('%Y%m%d_%H%M%S')}.json",
            "w",
            encoding="utf-8"
        ) as f:
            json.dump(data, f, indent=2, ensure_ascii=False)

        return data

    # ------------------------------------------------------------------
    # WORD TEMPLATE POPULATION
    # ------------------------------------------------------------------

    def _format_content(self, text):
        """Format text content with proper bullet points and numbered lists."""
        if not text:
            return ""

        lines = text.strip().split('\n')
        formatted_lines = []

        for line in lines:
            line = line.strip()
            if not line:
                continue

            # Check if line starts with numbered list pattern (e.g., "1.", "2.", etc.)
            if line and line[0].isdigit() and len(line) > 1 and line[1] == '.':
                formatted_lines.append(line)
            # Check if line starts with bullet point pattern (e.g., "- ", "* ", "• ")
            elif line.startswith(('- ', '* ', '• ')):
                formatted_lines.append(line)
            # For single lines or paragraphs, convert to bullet points
            elif '\n' not in line and not line.startswith(('- ', '* ', '• ')) and not (line and line[0].isdigit() and len(line) > 1 and line[1] == '.'):
                formatted_lines.append(f"- {line}")
            else:
                formatted_lines.append(line)

        return '\n'.join(formatted_lines)

    def fill_word_template(self, data: Dict, template_path: str, output_path: str) -> str:
        doc = Document(template_path)

       # ---------- DYNAMIC TABLES ----------
        try:
            br_table = self._find_table_by_placeholder(doc, "{req_id_bs}")
            if isinstance(data.get("business_requirements"), list):
                self._populate_business_requirements(doc, br_table, data["business_requirements"])
        except RuntimeError as e:
            print(f"WARNING: Business requirements table not found: {e}")

        try:
            tm_table = self._find_table_by_placeholder(doc, "{req_id_tm}")
            if isinstance(data.get("traceability_matrix"), list):
                self._populate_traceability(doc, tm_table, data["traceability_matrix"])
        except RuntimeError as e:
            print(f"WARNING: Traceability matrix table not found: {e}")

        try:
            try:
                toc_table = self._find_table_by_placeholder(doc, "{serial_number}")
            except RuntimeError:
                toc_table = self._find_table(doc, "Table Of Content")

            self._populate_table_of_contents(toc_table)
        except RuntimeError:
            print("WARNING: TOC table not found.")

        try:
            nfr_table = self._find_table_by_placeholder(doc, "{no_of_users}")
            self._populate_nfr(nfr_table, data)
        except RuntimeError as e:
            print(f"WARNING: NFR table not found: {e}")

        # ---------- SAFE PLACEHOLDER REPLACEMENT ----------
        flattened_data = self._get_flattened_data(data)

        def replace_text(text: str) -> str:
            for k, v in flattened_data.items():
                placeholder = f"{{{k}}}"
                if placeholder in text:
                    # Format the content if it's text content
                    formatted_value = self._format_content(str(v) if v is not None else "")
                    text = text.replace(placeholder, formatted_value)
            return text

        for p in doc.paragraphs:
            self._replace_text_in_runs(p, replace_text)

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        self._replace_text_in_runs(para, replace_text)

        # ---------- DIAGRAM INSERTION ----------
        # Check for Mermaid code in the LLM response
        visuals = data.get("process_visuals", {})
        mermaid_code = visuals.get("mermaid_code")
        if mermaid_code:
            success = self._insert_diagram(doc, mermaid_code)
            if not success:
                print("DEBUG: Diagram insertion failed or placeholder not found.")

        doc.save(output_path)
        return output_path

    # ------------------------------------------------------------------
    # FORMAT-SAFE RUN REPLACEMENT (FIX)
    # ------------------------------------------------------------------

    def _replace_text_in_runs(self, paragraph, replace_fn):
        for run in paragraph.runs:
            if run.text:
                run.text = replace_fn(run.text)

    # ------------------------------------------------------------------
    # TABLE HELPERS
    # ------------------------------------------------------------------

    def _find_table(self, doc, keyword: str):
        for table in doc.tables:
            if table.rows and keyword.lower() in table.rows[0].cells[0].text.lower():
                return table
        raise RuntimeError(f"Table not found: {keyword}")

    def _find_table_by_placeholder(self, doc, placeholder: str):
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if placeholder in cell.text:
                        return table
        raise RuntimeError(f"Table with placeholder '{placeholder}' not found")

    # --- NEW: Force visible borders on tables ---
    def _apply_table_borders(self, table):
        """
        Force borders on a python-docx table (outer + inner grid).
        """
        try:
            tbl = table._tbl
            tblPr = tbl.tblPr

            def set_border(tag, size="8", color="000000"):
                el = OxmlElement(tag)
                el.set(qn('w:val'), 'single')
                el.set(qn('w:sz'), size)   # 8 ≈ thin; increase for thicker lines
                el.set(qn('w:space'), '0')
                el.set(qn('w:color'), color)
                return el

            borders = OxmlElement('w:tblBorders')
            borders.append(set_border('w:top'))
            borders.append(set_border('w:bottom'))
            borders.append(set_border('w:left'))
            borders.append(set_border('w:right'))
            borders.append(set_border('w:insideH'))
            borders.append(set_border('w:insideV'))

            # Remove existing borders then apply
            for child in list(tblPr):
                if child.tag == qn('w:tblBorders'):
                    tblPr.remove(child)

            tblPr.append(borders)
        except Exception as e:
            print(f"DEBUG: Failed to apply table borders: {e}")

    def _duplicate_table_after(self, base_table, parent_doc):
        new_tbl_xml = deepcopy(base_table._tbl)
        parent = base_table._tbl.getparent()

        spacer_p = parent_doc.add_paragraph()
        spacer_p._p.getparent().remove(spacer_p._p)

        base_table._tbl.addnext(spacer_p._p)
        spacer_p._p.addnext(new_tbl_xml)

        new_table = Table(new_tbl_xml, base_table._parent)
        # Ensure duplicated tables have borders
        self._apply_table_borders(new_table)
        return new_table

    def _populate_business_requirements(self, doc, base_table, items):
        # Ensure the base table has borders
        self._apply_table_borders(base_table)

        current_table = base_table
        for idx, req in enumerate(items):
            target_table = base_table if idx == 0 else self._duplicate_table_after(current_table, doc)
            current_table = target_table

            # Also ensure borders on the (first/duplicated) target table
            self._apply_table_borders(target_table)

            target_table.cell(0, 1).text = str(req.get("req_id_bs", ""))
            target_table.cell(1, 1).text = str(req.get("title_bs", ""))
            target_table.cell(2, 1).text = self._format_content(str(req.get("description_bs", "")))
            target_table.cell(3, 1).text = self._format_content(str(req.get("as_is_behaviour", "")))
            target_table.cell(4, 1).text = self._format_content(str(req.get("to_be_behaviour", "")))
            target_table.cell(5, 1).text = self._format_content(str(req.get("pre_requisite", "")))
            target_table.cell(6, 1).text = self._format_content(str(req.get("acceptance_criteria", "")))
            target_table.cell(7, 1).text = self._format_content(str(req.get("alternate_flows", "")))
            target_table.cell(8, 1).text = self._format_content(str(req.get("reference_documents", "")))

    def _populate_traceability(self, doc, base_table, items):
        """
        Populates the single RTM table by adding rows for each item.
        The template row (with placeholders) is the second row (index 1).
        """
        for idx, trace_item in enumerate(items):
            # For the first item, use the existing placeholder row (index 1)
            # For subsequent items, add a new row at the bottom
            if idx == 0:
                row_cells = base_table.rows[1].cells
            else:
                row_cells = base_table.add_row().cells

            # Map the horizontal columns: REQ ID, Description, Source, Impacted System, Outcome
            # We use .strip() to keep content crisp as requested
            row_cells[0].text = str(trace_item.get("req_id_tm", "")).strip()
            row_cells[1].text = str(trace_item.get("description_tm", "")).strip()
            row_cells[2].text = str(trace_item.get("source_channel", "")).strip()
            row_cells[3].text = str(trace_item.get("impacted_system", "")).strip()
            row_cells[4].text = str(trace_item.get("outcome", "")).strip()

        # Ensure RTM table shows borders
        self._apply_table_borders(base_table)

    def _populate_table_of_contents(self, table):
        toc_items = [
            "Document Sign off", "Document History", "Overview", "Current constraints",
            "Objective", "In scope", "Out of scope", "Description",
            "Business Requirements", "Requirement Traceability Matrix",
            "Non-Functional Requirements", "Impact on Operational Process",
            "Regulatory Impact", "Reports Requirement", "Access Requirement",
            "Security Requirement", "Data Requirement", "Training Requirement"
        ]

        self._clear_table_keep_header(table)

        for i, section in enumerate(toc_items, 1):
            row = table.add_row().cells
            row[0].text = str(i)
            row[1].text = section

        # Ensure TOC table shows borders
        self._apply_table_borders(table)

    def _clear_table_keep_header(self, table):
        while len(table.rows) > 1:
            table._tbl.remove(table.rows[1]._tr)

    def _get_flattened_data(self, data: Dict) -> Dict:
        flat = {}
        if "document" in data:
            flat.update(data["document"])

        nfr = data.get("non_functional_requirements", {})
        if isinstance(nfr, list) and nfr:
            flat.update(nfr[0])

        for key in [
            "impact_on_operational_process", "regulatory_impact",
            "reports_requirement", "access_requirement",
            "security_requirement", "data_requirement",
            "training_requirement", "open_questions",
            "contradictions_found_in_input_documents"
        ]:
            if key in data:
                flat[key] = data[key]

        flat["date_today"] = datetime.today().strftime("%d-%b-%Y")
        return flat

    def _populate_nfr(self, table, data):
        nfr = data.get("non_functional_requirements", {})
        if isinstance(nfr, list) and nfr:
            nfr = nfr[0]

        table.cell(0, 1).text = self._format_content(str(nfr.get("no_of_users", "")))
        table.cell(1, 1).text = self._format_content(str(nfr.get("peak_volume", "")))
        table.cell(2, 1).text = self._format_content(str(nfr.get("monthly_volume", "")))
        table.cell(3, 1).text = self._format_content(str(nfr.get("availability", "")))

        # Ensure NFR table shows borders
        self._apply_table_borders(table)

    # ------------------------------------------------------------------
    # DIAGRAM LAYOUT HELPERS
    # ------------------------------------------------------------------

    def _set_row_cant_split(self, row):
        """
        Apply w:cantSplit to the table row to prevent Word from splitting it across pages.
        """
        try:
            tr = row._tr
            trPr = tr.get_or_add_trPr()
            if trPr.find(qn('w:cantSplit')) is None:
                cant_split = OxmlElement('w:cantSplit')
                trPr.append(cant_split)
        except Exception as e:
            print(f"DEBUG: Failed to set cantSplit on row: {e}")

    def _disable_row_splitting_all_tables(self, doc):
        """
        Apply cantSplit to all table rows in the document (defensive).
        """
        for table in doc.tables:
            for row in table.rows:
                self._set_row_cant_split(row)

    def _compute_width_for_page(self, doc, image_path):
        """
        Return a Length suitable for run.add_picture(width=...) that ensures:
        - Width does not exceed printable width
        - Height does not exceed printable height
        - Aspect ratio preserved
        """
        try:
            section = doc.sections[0]
            max_w_in = section.page_width.inches - section.left_margin.inches - section.right_margin.inches
            max_h_in = section.page_height.inches - section.top_margin.inches - section.bottom_margin.inches
        except Exception:
            # Safe defaults if sections are not accessible
            max_w_in = 6.0
            max_h_in = 8.0

        try:
            with Image.open(image_path) as im:
                w_px, h_px = im.size
        except Exception as e:
            print(f"DEBUG: PIL failed to open image for size calc: {e}")
            safe_w = min(5.5, max_w_in)
            return Inches(safe_w)

        if w_px <= 0 or h_px <= 0:
            safe_w = min(5.5, max_w_in)
            return Inches(safe_w)

        aspect = h_px / float(w_px)

        # Start with max printable width
        width_in = max_w_in
        height_in = width_in * aspect

        # If height would exceed printable height, reduce width accordingly
        if height_in > max_h_in:
            height_in = max_h_in
            width_in = height_in / aspect

        # Tiny safety cushion to avoid Word rounding overflow at edges
        width_in = max(0.5, min(width_in, max_w_in) - 0.02)
        return Inches(width_in)

    # ------------------------------------------------------------------
    # DIAGRAM GENERATION & INSERTION (HARDENED)
    # ------------------------------------------------------------------
    def _insert_diagram(self, doc, mermaid_code, placeholder="{process_diagram}"):
        """
        Convert Mermaid code to PNG using Mermaid.ink API and insert into document.
        Ensures:
        - Image is inserted in a NEW paragraph (not inline)
        - Placeholder paragraph is removed
        - Title paragraph remains untouched
        - Correct order: Title -> Image
        - No pagination glitches
        - No inline layout bugs
        - Professional BRD formatting
        """

        try:
            # -----------------------------
            # 1) Clean Mermaid code
            # -----------------------------
            mermaid_code = (mermaid_code or "").replace("```mermaid", "").replace("```", "").strip()
            mermaid_code = mermaid_code.replace('\xa0', ' ')
            if not mermaid_code:
                print("DEBUG: Mermaid code is empty; skipping diagram insertion.")
                return False

            # -----------------------------
            # 2) Encode + fetch image
            # -----------------------------
            graphbytes = mermaid_code.encode("utf-8")
            base64_string = base64.urlsafe_b64encode(graphbytes).decode("ascii").rstrip("=")
            print(f"Base64 string: {base64_string}")
            url = f"https://mermaid.ink/img/{base64_string}"

            response = None
            for attempt in range(3):
                try:
                    response = requests.get(url, timeout=30)
                    if response.status_code == 200:
                        break
                    print(f"DEBUG: Mermaid API status {response.status_code}, attempt {attempt+1}/3")
                except Exception as e:
                    print(f"DEBUG: Mermaid API request error (attempt {attempt+1}/3): {e}")
                time.sleep(1)

            if not response or response.status_code != 200:
                print("DEBUG: Mermaid API failed; diagram not generated.")
                return False

            image_path = "process_flow.png"
            with open(image_path, "wb") as f:
                f.write(response.content)

            # -----------------------------
            # 3) Compute safe width
            # -----------------------------
            width_len = self._compute_width_for_page(doc, image_path)

            # -----------------------------
            # 4) Find placeholder
            # -----------------------------
            placeholder_para = None
            placeholder_parent = None

            # Search body paragraphs
            for p in doc.paragraphs:
                if placeholder in p.text:
                    placeholder_para = p
                    placeholder_parent = p._p.getparent()
                    break

            # Search tables if not found
            if not placeholder_para:
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for p in cell.paragraphs:
                                if placeholder in p.text:
                                    placeholder_para = p
                                    placeholder_parent = p._p.getparent()
                                    # prevent row split
                                    self._set_row_cant_split(row)
                                    break
                            if placeholder_para:
                                break
                        if placeholder_para:
                            break
                    if placeholder_para:
                        break

            if not placeholder_para:
                print(f"DEBUG: Placeholder '{placeholder}' not found.")
                return False

            # -----------------------------
            # 5) Remove placeholder paragraph
            # -----------------------------
            # Remove only the placeholder text, not the whole paragraph
            for r in placeholder_para.runs:
                if r.text and placeholder in r.text:
                    r.text = r.text.replace(placeholder, "").strip()

            # If paragraph becomes empty after removal, then delete it
            if not placeholder_para.text.strip():
                placeholder_parent.remove(placeholder_para._p)

            # -----------------------------
            # 6) Insert image in NEW paragraph (BLOCK)
            # -----------------------------
            img_para = doc.add_paragraph()
            img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            img_para.paragraph_format.keep_together = True
            img_para.paragraph_format.keep_with_next = False
            img_para.paragraph_format.space_before = Pt(6)
            img_para.paragraph_format.space_after = Pt(12)

            run = img_para.add_run()
            run.add_picture(image_path, width=width_len)

            return True

        except Exception as e:
            print(f"DEBUG: Error in _insert_diagram: {e}")
            return False

    # ------------------------------------------------------------------
    # ORCHESTRATOR
    # ------------------------------------------------------------------

    def generate_brd_word(
        self,
        requirements: str,
        context: str,
        schema_json: str,
        template_path: str,
        output_path: str,
        process_instructions: str = ""
    ) -> str:
        data = self.generate_brd_json(requirements, context, schema_json, process_instructions)
        return self.fill_word_template(data, template_path, output_path)